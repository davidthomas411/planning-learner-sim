from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "paper" / "planning_trajectory_manuscript_draft.md"
OUTPUT = ROOT / "paper" / "planning_trajectory_manuscript_draft.docx"
QA_NO_LINE_NUMBERS = os.environ.get("DOCX_QA_NO_LINE_NUMBERS") == "1"
QA_SIMPLE = os.environ.get("DOCX_QA_SIMPLE") == "1"
if QA_NO_LINE_NUMBERS:
    suffix = "simple" if QA_SIMPLE else "qa"
    OUTPUT = ROOT / ".docx_review" / "manuscript_v1" / f"planning_trajectory_manuscript_{suffix}.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, end))


def configure_section(section, line_numbers: bool) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    if line_numbers:
        sect_pr = section._sectPr
        old = sect_pr.find(qn("w:lnNumType"))
        if old is not None:
            sect_pr.remove(old)
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:distance"), "360")
        ln.set(qn("w:restart"), "continuous")
        sect_pr.append(ln)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Inches(0.25)
    normal.paragraph_format.widow_control = True

    for name, size, before, after in (
        ("Title", 16, 0, 14),
        ("Heading 1", 13, 14, 6),
        ("Heading 2", 11, 10, 3),
        ("Heading 3", 11, 8, 2),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 56, 92) if name != "Title" else RGBColor(18, 45, 78)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Inches(0)
        style.paragraph_format.line_spacing = 1.0

    if "Caption Scientific" not in styles:
        cap = styles.add_style("Caption Scientific", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Caption Scientific"]
    cap.font.name = "Arial"
    cap.font.size = Pt(9)
    cap.font.italic = False
    cap.paragraph_format.first_line_indent = Inches(0)
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_with_next = True

    if "Table Text" not in styles:
        table_style = styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_style = styles["Table Text"]
    table_style.font.name = "Arial"
    table_style.font.size = Pt(8.5)
    table_style.paragraph_format.first_line_indent = Inches(0)
    table_style.paragraph_format.line_spacing = 1.0
    table_style.paragraph_format.space_after = Pt(0)


def add_inline_markup(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    usable = Inches(6.5)
    widths = [usable / cols for _ in range(cols)]
    if cols == 2:
        widths = [Inches(1.65), Inches(4.85)]
    elif cols == 3:
        widths = [Inches(1.35), Inches(2.55), Inches(2.60)]
    elif cols == 4:
        widths = [Inches(1.45), Inches(1.65), Inches(1.7), Inches(1.7)]
    elif cols == 5:
        widths = [Inches(1.55), Inches(1.15), Inches(1.35), Inches(1.55), Inches(0.9)]
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        row.cantSplit = True
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, cell in enumerate(row.cells):
            cell.width = widths[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "D9E5F3")
            value = row_data[c_idx] if c_idx < len(row_data) else ""
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Text"]
            add_inline_markup(p, value)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(doc: Document, image_path: Path) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(image_path), width=Inches(6.25))


def style_placeholder(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "FFF2CC")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), "BF9000")
        borders.append(edge)
    p_pr.append(borders)


def add_header_footer(section, show_header: bool) -> None:
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "PLANNING TRAJECTORY SUPERVISION" if show_header else ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(89, 89, 89)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if show_header:
        add_field(fp, "PAGE")
    for run in fp.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)


def build() -> None:
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], line_numbers=False)
    add_header_footer(doc.sections[0], show_header=False)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    title_seen = False
    in_references = False
    pending_table_label = False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line == "\\pagebreak":
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            configure_section(section, line_numbers=not QA_NO_LINE_NUMBERS)
            add_header_footer(section, show_header=True)
            i += 1
            continue
        if line.startswith("# "):
            text = line[2:].strip()
            if not title_seen:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(36)
                add_inline_markup(p, text)
                title_seen = True
            else:
                p = doc.add_paragraph(text, style="Heading 1")
                in_references = text == "References"
            i += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
            i += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 3")
            i += 1
            continue
        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            image_path = (SOURCE.parent / image_match.group(2)).resolve()
            if not QA_SIMPLE:
                add_image(doc, image_path)
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            if not QA_SIMPLE:
                add_table(doc, rows)
            pending_table_label = False
            continue
        if re.match(r"^\d+\.\s", line) and not in_references:
            p = doc.add_paragraph(style="Normal")
            p.style = doc.styles["Normal"]
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.hanging_indent = Inches(0.25)
            add_inline_markup(p, line)
            i += 1
            continue
        p = doc.add_paragraph(style="Normal")
        add_inline_markup(p, line)
        if line.startswith("**Figure"):
            p.style = doc.styles["Caption Scientific"]
        elif line.startswith("**Table"):
            p.style = doc.styles["Caption Scientific"]
            p.paragraph_format.keep_with_next = True
            pending_table_label = True
        elif line.startswith("[RESULTS PLACEHOLDER") or line.startswith("[CONCLUSION PLACEHOLDER"):
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.right_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            style_placeholder(p)
        elif in_references and re.match(r"^\d+\.", line):
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(4)
        elif not title_seen or len(doc.sections) == 1:
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
        i += 1

    props = doc.core_properties
    props.title = "Do planning trajectories provide information beyond final plans?"
    props.subject = "Protocol-stage Medical Physics research article draft"
    props.author = "David Thomas"
    props.keywords = "radiation treatment planning; trajectory supervision; simulation"
    props.comments = "Results are intentionally reserved as prespecified placeholders."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
